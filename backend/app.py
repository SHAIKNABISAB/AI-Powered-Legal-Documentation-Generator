import os
import io
import datetime
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

app = Flask(__name__)
CORS(app)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = "gemini-1.5-flash"
MAX_ANALYSIS_TEXT_LENGTH = 8000

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

ALLOWED_UPLOAD_EXTENSIONS = {"pdf", "docx", "txt", "doc"}
MAX_UPLOAD_SIZE = 10 * 1024 * 1024  # 10 MB

DOCUMENT_TYPES = {
    "employment_contract": {
        "title": "Employment Contract",
        "fields": ["employer_name", "employee_name", "position", "start_date", "salary", "work_location", "employment_type"],
    },
    "lease_agreement": {
        "title": "Lease Agreement",
        "fields": ["landlord_name", "tenant_name", "property_address", "lease_start", "lease_end", "monthly_rent", "security_deposit"],
    },
    "nda": {
        "title": "Non-Disclosure Agreement",
        "fields": ["disclosing_party", "receiving_party", "effective_date", "jurisdiction", "term_years"],
    },
    "partnership_agreement": {
        "title": "Partnership Agreement",
        "fields": ["partner_one", "partner_two", "business_name", "business_address", "start_date", "profit_sharing", "duration"],
    },
    "service_agreement": {
        "title": "Service Agreement",
        "fields": ["service_provider", "client_name", "service_description", "start_date", "end_date", "payment_amount", "payment_terms"],
    },
    "settlement_agreement": {
        "title": "Settlement Agreement",
        "fields": ["party_one", "party_two", "dispute_description", "settlement_amount", "settlement_date", "confidentiality_required"],
    },
    "freelance_contract": {
        "title": "Freelance Contract",
        "fields": ["freelancer_name", "client_name", "project_description", "start_date", "deadline", "total_payment", "milestone_payment"],
    },
    "terms_of_service": {
        "title": "Terms of Service",
        "fields": ["company_name", "website_url", "effective_date", "governing_law", "minimum_age"],
    },
    "domicile_agreement": {
        "title": "Domicile Agreement",
        "fields": ["resident_name", "property_owner", "property_address", "start_date", "monthly_fee", "services_included"],
    },
}


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_UPLOAD_EXTENSIONS


def extract_text_from_file(file_storage):
    filename = file_storage.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    content = file_storage.read()

    if ext == "txt":
        return content.decode("utf-8", errors="replace")

    if ext in ("docx", "doc"):
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    if ext == "pdf":
        try:
            import fitz  # PyMuPDF
            pdf_doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            for page in pdf_doc:
                text_parts.append(page.get_text())
            pdf_doc.close()
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Could not extract text from PDF: {e}")

    raise ValueError(f"Unsupported file type: {ext}")


def build_docx(document_text, document_title):
    doc = Document()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(document_title.upper())
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_paragraph()

    for line in document_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("**") and stripped.endswith("**"):
            heading_text = stripped.strip("*").strip()
            para = doc.add_paragraph()
            run = para.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(12)
        else:
            cleaned = stripped.replace("**", "")
            para = doc.add_paragraph(cleaned)
            para.style.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.route("/", methods=["GET"])
def index():
    return jsonify({
        "message": "LegalEase API is running",
        "version": "1.0.0",
        "endpoints": ["/document-types", "/generate", "/download-doc", "/upload"],
    })


@app.route("/document-types", methods=["GET"])
def get_document_types():
    return jsonify({
        "success": True,
        "document_types": {k: {"title": v["title"], "fields": v["fields"]} for k, v in DOCUMENT_TYPES.items()},
    })


@app.route("/generate", methods=["POST"])
def generate_document():
    data = request.get_json()
    if not data:
        return jsonify({"success": False, "error": "No data provided"}), 400

    document_type = data.get("document_type", "").strip()
    details = data.get("details", {})

    if not document_type or document_type not in DOCUMENT_TYPES:
        return jsonify({"success": False, "error": "Invalid document type"}), 400

    doc_info = DOCUMENT_TYPES[document_type]
    details_text = "\n".join(f"- {k.replace('_', ' ').title()}: {v}" for k, v in details.items())

    prompt = (
        f"Generate a professional, legally structured {doc_info['title']} document.\n\n"
        f"Document Details:\n{details_text}\n\n"
        "Requirements:\n"
        "1. Use formal legal language\n"
        "2. Include all standard clauses for this document type\n"
        "3. Structure with clear numbered sections\n"
        "4. Use **bold** for section headings\n"
        "5. Include a signature section at the end\n"
        "6. Make it comprehensive and professional\n\n"
        "Generate the complete document now:"
    )

    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        document_text = response.text
    except Exception as e:
        return jsonify({"success": False, "error": f"AI generation failed: {str(e)}"}), 500

    return jsonify({
        "success": True,
        "document_type": document_type,
        "document_title": doc_info["title"],
        "document_text": document_text,
        "generated_at": datetime.datetime.now().isoformat(),
        "message": "Document generated successfully",
    })


@app.route("/download-doc", methods=["POST"])
def download_doc():
    data = request.get_json()
    if not data:
        return jsonify({"success": False, "error": "No data provided"}), 400

    document_text = data.get("document_text", "").strip()
    document_type = data.get("document_type", "document")

    if not document_text:
        return jsonify({"success": False, "error": "No document text provided"}), 400

    doc_info = DOCUMENT_TYPES.get(document_type, {"title": "Legal Document"})
    document_title = doc_info["title"] if isinstance(doc_info, dict) else "Legal Document"

    buffer = build_docx(document_text, document_title)
    date_str = datetime.datetime.now().strftime("%Y%m%d")
    filename = f"LegalEase_{document_title.replace(' ', '_')}_{date_str}.docx"

    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return jsonify({"success": False, "error": "No file provided. Include a file in the 'file' field."}), 400

    file = request.files["file"]

    if not file.filename:
        return jsonify({"success": False, "error": "No file selected."}), 400

    if not allowed_file(file.filename):
        return jsonify({
            "success": False,
            "error": f"File type not supported. Allowed types: {', '.join(sorted(ALLOWED_UPLOAD_EXTENSIONS))}",
        }), 400

    file.seek(0, 2)
    size = file.tell()
    file.seek(0)
    if size > MAX_UPLOAD_SIZE:
        return jsonify({"success": False, "error": "File exceeds maximum allowed size of 10 MB."}), 400

    try:
        extracted_text = extract_text_from_file(file)
    except ValueError as e:
        return jsonify({"success": False, "error": str(e)}), 422
    except Exception as e:
        return jsonify({"success": False, "error": f"Failed to process file: {str(e)}"}), 500

    if not extracted_text.strip():
        return jsonify({"success": False, "error": "No readable text found in the uploaded file."}), 422

    summary_prompt = (
        "You are a legal document analyst. Analyze the following document and provide:\n"
        "1. Document type (e.g., Employment Contract, Lease Agreement, etc.)\n"
        "2. Key parties involved\n"
        "3. Main terms and conditions (bullet points)\n"
        "4. Important dates and deadlines\n"
        "5. Any notable clauses or provisions\n\n"
        f"Document content:\n{extracted_text[:MAX_ANALYSIS_TEXT_LENGTH]}\n\n"
        "Provide a structured analysis:"
    )

    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        response = model.generate_content(summary_prompt)
        analysis = response.text
    except Exception as e:
        analysis = None

    return jsonify({
        "success": True,
        "filename": file.filename,
        "extracted_text": extracted_text,
        "analysis": analysis,
        "message": "File uploaded and processed successfully",
    })


if __name__ == "__main__":
    debug = os.getenv("FLASK_DEBUG", "false").lower() == "true"
    app.run(debug=debug, host="0.0.0.0", port=5000)
